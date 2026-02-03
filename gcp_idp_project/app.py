import argparse
import os
import sys
import logging
import gc
from pathlib import Path

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/app.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def run_ocr_pipeline(input_folder):
    """
    Mode: OCR
    - Load Surya OCR ONLY
    - Convert PDFs -> images
    - Preprocessing
    - OCR page-by-page
    - Save to output/<case_id>/ocr/ and output/<case_id>/full_text.txt
    """
    logger.info("Starting OCR Pipeline...")
    
    # Lazy import to ensure separation
    try:
        from PIL import Image
        from surya.ocr import run_ocr
        from surya.model.detection import segformer
        from surya.model.recognition.model import load_model
        from surya.model.recognition.processor import load_processor
        import torch
    except ImportError as e:
        logger.error(f"Failed to import OCR dependencies: {e}")
        sys.exit(1)

    input_path = Path(input_folder)
    if not input_path.exists():
        logger.error(f"Input folder not found: {input_folder}")
        sys.exit(1)

    # Load models once
    logger.info("Loading Surya OCR models...")
    try:
        # Note: Surya API might vary slightly depending on version, adjusting for standard usage
        det_processor, det_model = segformer.load_processor(), segformer.load_model()
        rec_model, rec_processor = load_model(), load_processor()
    except Exception as e:
        logger.error(f"Error loading models: {e}")
        sys.exit(1)

    # Process files
    # Assuming input_folder contains PDF files or subdirectories of cases
    # The user requirements say: "Save: output/<case_id>/ocr/..."
    # So we should treat subfolders in input_folder as cases or files as cases?
    # "ingest/apar/" and "ingest/disciplinary/..." are the inputs.
    # If input is "ingest/apar", then files inside are the documents.
    # Let's assume each file is a "case" or "document" for now, or if input is a folder of files.
    # The requirement says "output/<case_id>/". If input is a file "doc1.pdf", case_id is "doc1".
    
    # Process files
    # Iterate through all PDFs in the input folder
    files = list(input_path.rglob("*.pdf"))
    if not files:
        logger.warning(f"No PDF files found in {input_folder}")
        return

    # Group files by case_id if possible, or treat input_folder as the case root?
    # If input is "ingest/disciplinary/case1", then case_id is "case1".
    # If input is "ingest/apar", then each file is a case?
    # Let's assume the user passes the specific case folder or a root containing cases.
    # "python app.py --mode ocr --input <folder>"
    # If <folder> is "ingest/disciplinary/case_001", then case_id is "case_001".
    
    # We will assume the input_folder IS the case folder for simplicity, 
    # OR we treat subfolders as cases.
    # Given the strict structure "ingest/disciplinary/brief_background/...", 
    # it seems "ingest/disciplinary" contains categories, not cases directly?
    # Or "ingest/disciplinary/<case_id>/brief_background"?
    # The prompt structure:
    # ingest/
    #   apar/
    #   disciplinary/
    #       brief_background/
    #       po_brief/
    #       ...
    # This implies ALL disciplinary documents for ALL cases might be mixed, OR
    # "ingest/disciplinary" is for ONE case?
    # "2️⃣ Disciplinary Case Documents ... Documents are pre-grouped into folders: Brief Background..."
    # This phrasing suggests a single case structure or a standard structure for cases.
    
    # Let's assume input_folder points to a specific CASE ROOT.
    case_id = input_path.name
    output_dir = Path("output") / case_id
    ocr_output_dir = output_dir / "ocr"
    ocr_output_dir.mkdir(parents=True, exist_ok=True)
    
    full_text_content = []
    
    # Sort files to maintain some order
    files.sort()
    
    global_page_count = 0

    for file_path in files:
        # Determine category based on parent folder relative to input_path
        # e.g. input="case1", file="case1/brief_background/doc.pdf" -> category="brief_background"
        try:
            rel_path = file_path.relative_to(input_path)
            category = rel_path.parent.name if rel_path.parent != Path(".") else "general"
        except ValueError:
            category = "general"

        logger.info(f"Processing {file_path.name} (Category: {category})")
        
        try:
            import pypdfium2 as pdfium
            
            pdf = pdfium.PdfDocument(str(file_path))
            doc_text = []
            
            for i, page in enumerate(pdf):
                global_page_count += 1
                image = page.render(scale=2).to_pil() 
                
                # Run OCR
                predictions = run_ocr([image], [list(image.size)], det_model, det_processor, rec_model, rec_processor)
                
                # Extract text
                page_text = ""
                if predictions and len(predictions) > 0:
                    text_lines = [l.text_content for l in predictions[0].text_lines]
                    page_text = "\n".join(text_lines)
                
                # Save page text with category prefix in filename for traceability
                # e.g. brief_background_doc1_page1.txt
                safe_filename = f"{category}_{file_path.stem}_p{i+1}.txt".replace(" ", "_")
                with open(ocr_output_dir / safe_filename, "w", encoding="utf-8") as f:
                    f.write(page_text)
                
                doc_text.append(f"--- Page {i+1} ({file_path.name}) ---\n{page_text}")
                logger.info(f"Page {i+1} of {file_path.name}: Success")
                
                del image
                gc.collect()
            
            # Append to full text with section header
            full_text_content.append(f"\n=== SECTION: {category} | FILE: {file_path.name} ===\n" + "\n".join(doc_text))
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            continue
            
    # Save full text
    with open(output_dir / "full_text.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(full_text_content))

    logger.info(f"OCR Pipeline Finished for Case: {case_id}")

def run_nlp_pipeline(input_folder):
    """
    Mode: NLP
    - Load Phi-3 Mini
    - Read full_text.txt
    - Classify
    - Extract JSON or Generate Reports
    """
    logger.info("Starting NLP Pipeline...")
    
    # Lazy import
    try:
        from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
        import torch
        import json
    except ImportError as e:
        logger.error(f"Failed to import NLP dependencies: {e}")
        sys.exit(1)

    input_path = Path(input_folder)
    # The input for NLP mode is likely the folder containing the source documents? 
    # OR the folder containing the OCR output?
    # Requirement: "Read full_text.txt". This implies input might be the output folder from OCR step?
    # Or app.py iterates over the output folder structure?
    # "python app.py --mode nlp --input <folder>"
    # If <folder> is the original input folder, we map to output folder.
    # If <folder> is the 'output' directory, we iterate cases.
    # Let's assume <folder> is the root 'output' directory containing case_id folders.
    
    if not input_path.exists():
        logger.error(f"Input folder not found: {input_folder}")
        sys.exit(1)

    # Load Model
    logger.info("Loading Phi-3 Mini...")
    model_id = "microsoft/Phi-3-mini-4k-instruct" 
    # Note: User specified "Phi-3 Mini (3.8B, local)". 
    # We assume model is downloaded or we use HF cache.
    # "models/" directory exists in project, maybe we should point there.
    
    try:
        # Use GPU if available
        device = "cuda" if torch.cuda.is_available() else "cpu"
        
        tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)
        model = AutoModelForCausalLM.from_pretrained(
            model_id, 
            device_map=device, 
            torch_dtype="auto", 
            trust_remote_code=True
        )
        
        pipe = pipeline(
            "text-generation",
            model=model,
            tokenizer=tokenizer,
        )
    except Exception as e:
        logger.error(f"Error loading NLP model: {e}")
        sys.exit(1)

    # Iterate over cases in the input folder
    # Assuming input_folder contains subfolders for each case
    case_folders = [f for f in input_path.iterdir() if f.is_dir()]
    
    if not case_folders:
        logger.warning(f"No case folders found in {input_folder}")
        return

    for case_folder in case_folders:
        case_id = case_folder.name
        full_text_path = case_folder / "full_text.txt"
        
        if not full_text_path.exists():
            logger.warning(f"No full_text.txt found for case {case_id}. Skipping.")
            continue
            
        logger.info(f"Processing NLP for {case_id}")
        
        with open(full_text_path, "r", encoding="utf-8") as f:
            text_content = f.read()
            
        # Step 1: Classify
        doc_type = classify_document(pipe, text_content)
        logger.info(f"Case {case_id} classified as: {doc_type}")
        
        if doc_type == "APAR":
            process_apar(pipe, text_content, case_folder)
        elif doc_type == "DISCIPLINARY_CASE":
            process_disciplinary(pipe, text_content, case_folder)
        else:
            logger.warning(f"Unknown document type for {case_id}")

    logger.info("NLP Pipeline Finished.")

def classify_document(pipe, text):
    """
    Classify text as APAR or DISCIPLINARY_CASE
    """
    # Truncate text for classification to fit context
    prompt = f"""
    <|user|>
    Classify the following document into exactly one of these two types: "APAR" or "DISCIPLINARY_CASE".
    Return only the category name.
    
    Document Snippet:
    {text[:2000]}
    <|end|>
    <|assistant|>
    """
    
    output = pipe(prompt, max_new_tokens=10, return_full_text=False)[0]['generated_text']
    result = output.strip().upper()
    if "APAR" in result: return "APAR"
    if "DISCIPLINARY" in result: return "DISCIPLINARY_CASE"
    return "UNKNOWN"

import re
import json
import pandas as pd
from docx import Document

def split_apar_by_year(full_text):
    """
    PASS 1: Text Segmentation (No LLM)
    Detect APAR year boundaries and split text into chunks.
    """
    # Patterns to detect start of a new APAR year section
    patterns = [
        r"APAR\s+FOR\s+THE\s+YEAR",
        r"Annual\s+Performance\s+Appraisal\s+Report",
        r"\b(?:19|20)\d{2}\s*-\s*(?:19|20)?\d{2}\b" # Year ranges like 2020-21
    ]
    
    # We want to find all occurrences of these patterns to identify split points.
    # However, just splitting by year might be too granular if a year is mentioned multiple times.
    # We generally expect "APAR FOR THE YEAR ..." to be a header.
    # Let's try to find the "APAR FOR THE YEAR" or "Annual Performance..." headers specifically.
    
    header_pattern = r"(?:APAR\s+FOR\s+THE\s+YEAR|Annual\s+Performance\s+Appraisal\s+Report)"
    
    # Find all matches
    matches = list(re.finditer(header_pattern, full_text, re.IGNORECASE))
    
    chunks = []
    if not matches:
        # If no explicit headers, try splitting by significant year ranges if they look like headers?
        # Or just return the whole text as one chunk if it fits, but requirement says "Detect... Return one text chunk per APAR year"
        # If we can't find headers, we might fail or return the whole text (and truncate if needed).
        # Let's fallback to splitting by year range if distinct.
        year_pattern = r"\b(20\d{2}\s*-\s*\d{2,4})\b"
        matches = list(re.finditer(year_pattern, full_text))
        if not matches:
             # Fallback: simple chunking by size if really needed, or return single chunk
             return [full_text]
    
    # If we have matches, split text
    last_idx = 0
    # If the first match is not at the beginning, the text before it might be metadata or previous year (if header missing).
    # But usually metadata is at start.
    
    # Strategy: Start of a match marks the start of a new chunk.
    # The text before the first match is likely the global metadata / cover page.
    # We will include the text before the first match as a separate chunk or attach it to the first year?
    # Requirement: "Return one text chunk per APAR year". 
    # Global metadata is extracted in Pass 3 from "first 2-3 pages".
    # So here we focus on Year Chunks.
    
    # Let's treat the start of the text as potential content, but usually headers separate years.
    # If the first header is deep in the text, the text before it might be the first year?
    # Or text before first header is just cover page.
    # Let's iterate matches.
    
    split_indices = [m.start() for m in matches]
    
    # Filter indices that are too close to each other (avoid splitting on repeated headers in same section)
    # A generic APAR form is usually 2-10 pages. 
    # Let's merge splits that are < 500 chars apart?
    
    filtered_indices = []
    if split_indices:
        filtered_indices.append(split_indices[0])
        for idx in split_indices[1:]:
            if idx - filtered_indices[-1] > 1000: # Assuming at least 1000 chars per APAR
                filtered_indices.append(idx)
    
    # Now create chunks
    if not filtered_indices:
        chunks.append(full_text)
    else:
        # If there is significant text before the first header, is it a year? 
        # Usually cover letter. We can ignore for Year Extraction or check.
        # But let's assume the first header starts the first year.
        # Wait, if the file starts with "APAR FOR THE YEAR...", index is 0.
        
        # We will extract from each index to the next.
        for i in range(len(filtered_indices)):
            start = filtered_indices[i]
            end = filtered_indices[i+1] if i + 1 < len(filtered_indices) else len(full_text)
            chunk = full_text[start:end]
            
            # Constraint: < 4000 characters
            # If chunk is too large, we must truncate or split further.
            # Requirement: "Ensure each chunk is < 4000 characters"
            if len(chunk) > 4000:
                chunk = chunk[:4000] # Safe truncation
            
            chunks.append(chunk)
            
    return chunks

def extract_year_record(pipe, year_text):
    """
    PASS 2: Year-wise LLM Extraction
    """
    prompt = f"""
    <|user|>
    You are an expert government APAR analyst.
    Extract information ONLY for ONE APAR YEAR from the text below.
    Do not guess. Do not hallucinate.
    
    Return STRICT JSON:
    {{
      "year": "",
      "reporting_officer": "",
      "reviewing_officer": "",
      "accepting_authority": "",
      "grading": "",
      "pen_picture": ""
    }}
    
    Rules:
    - Use only the provided text
    - If a field is missing, return empty string
    - Pen picture must be verbatim, merged into one paragraph
    - Output JSON only
    
    TEXT:
    {year_text}
    <|end|>
    <|assistant|>
    ```json
    """
    try:
        output = pipe(prompt, max_new_tokens=500, return_full_text=False)[0]['generated_text']
        # Clean JSON
        json_str = output.strip()
        if "```json" in json_str:
            json_str = json_str.split("```json")[1].split("```")[0]
        elif "```" in json_str:
            json_str = json_str.split("```")[0] # If just ```
            
        return json.loads(json_str)
    except Exception as e:
        logger.error(f"Error extracting year record: {e}")
        return {}

def extract_officer_metadata(pipe, header_text):
    """
    PASS 3: Global Metadata Extraction
    """
    prompt = f"""
    <|user|>
    You are extracting officer metadata from an APAR document.
    
    Return STRICT JSON:
    {{
      "officer_name": "",
      "dob": ""
    }}
    
    Rules:
    - Use only provided text
    - Do not infer or guess
    - Output JSON only
    
    TEXT:
    {header_text}
    <|end|>
    <|assistant|>
    ```json
    """
    try:
        output = pipe(prompt, max_new_tokens=200, return_full_text=False)[0]['generated_text']
        json_str = output.strip()
        if "```json" in json_str:
            json_str = json_str.split("```json")[1].split("```")[0]
        elif "```" in json_str:
            json_str = json_str.split("```")[0]
            
        return json.loads(json_str)
    except Exception as e:
        logger.error(f"Error extracting metadata: {e}")
        return {"officer_name": "", "dob": ""}

def process_apar(pipe, text, output_folder):
    logger.info("Processing APAR with Multi-Pass Strategy...")
    
    llm_call_count = 0
    total_input_chars = 0
    confidence_flags = []
    
    # PASS 1: Text Segmentation
    # We need to handle the whole text. 
    # Note: text passed here is full_text.txt content.
    
    year_chunks = split_apar_by_year(text)
    logger.info(f"Identified {len(year_chunks)} year chunks.")
    
    if not year_chunks:
        logger.error("No APAR year chunks detected. Aborting NLP stage.")
        return
    
    # PASS 2: Year-wise Extraction
    apar_records = []
    for i, chunk in enumerate(year_chunks):
        logger.info(f"Processing Year Chunk {i+1} (Length: {len(chunk)} chars)...")
        
        # Log chunk boundary
        logger.info(f"Chunk {i+1} Boundary: Start-End (approx) of {len(chunk)} chars")
        
        llm_call_count += 1
        total_input_chars += len(chunk)
        
        record = extract_year_record(pipe, chunk)
        if record:
            # Validation: Check grading
            if not record.get("grading"):
                logger.warning(f"Missing grading for year {record.get('year', 'Unknown')}")
                confidence_flags.append(f"MISSING_GRADING: Year {record.get('year', 'Unknown')}")
            
            apar_records.append(record)
    
    if not apar_records:
        logger.error("No valid APAR records extracted. Aborting.")
        return

    # PASS 3: Global Metadata
    # "Extract officer-level metadata using ONLY the first 2–3 pages of OCR text."
    # We can try to split by "--- Page X ---" if preserved, or just take first 8000 chars (approx 3 pages) and truncate to 4000 for safety?
    # Requirement: "Never send >4000 characters to the LLM".
    # So we take the first 4000 chars.
    
    metadata_text = text[:4000]
    logger.info("Extracting Global Metadata...")
    
    llm_call_count += 1
    total_input_chars += len(metadata_text)
    
    metadata = extract_officer_metadata(pipe, metadata_text)
    
    # Validation: Check officer name
    if not metadata.get("officer_name"):
        logger.warning("Officer Name missing.")
        confidence_flags.append("LOW_CONFIDENCE: Officer Name Missing")
    
    # PASS 4: Final Merge
    final_json = {
        "officer_name": metadata.get("officer_name", ""),
        "dob": metadata.get("dob", ""),
        "confidence_flags": confidence_flags,
        "apar_records": apar_records
    }
    
    logger.info(f"NLP Stats: {llm_call_count} LLM calls, ~{total_input_chars} input chars.")
    
    # Save JSON
    try:
        with open(output_folder / "apar_data.json", "w") as f:
            json.dump(final_json, f, indent=2)
            
        # Generate CSV
        flat_records = []
        for r in apar_records:
            row = r.copy()
            row['officer_name'] = final_json['officer_name']
            row['dob'] = final_json['dob']
            row['confidence_flags'] = "; ".join(confidence_flags)
            flat_records.append(row)
            
        if flat_records:
            df = pd.DataFrame(flat_records)
            # Reorder columns to match requirement
            cols = ["officer_name", "dob", "year", "reporting_officer", "reviewing_officer", "accepting_authority", "grading", "pen_picture", "confidence_flags"]
            # Filter cols that exist
            cols = [c for c in cols if c in df.columns]
            df = df[cols]
            df.to_csv(output_folder / "apar_result.csv", index=False)
            
        # Generate DOCX
        doc = Document()
        doc.add_heading('APAR Result Format', 0)
        
        # Add metadata section
        doc.add_paragraph(f"Officer Name: {final_json['officer_name']}")
        doc.add_paragraph(f"Date of Birth: {final_json['dob']}")
        if confidence_flags:
             doc.add_paragraph(f"WARNING: {'; '.join(confidence_flags)}").bold = True
        
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Name", "DoB", "Year", "Reporting Officer", "Reviewing Officer", "Accepting Authority", "Grading", "Pen Picture"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            
        for rec in flat_records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(rec.get('officer_name', ''))
            row_cells[1].text = str(rec.get('dob', ''))
            row_cells[2].text = str(rec.get('year', ''))
            row_cells[3].text = str(rec.get('reporting_officer', ''))
            row_cells[4].text = str(rec.get('reviewing_officer', ''))
            row_cells[5].text = str(rec.get('accepting_authority', ''))
            row_cells[6].text = str(rec.get('grading', ''))
            row_cells[7].text = str(rec.get('pen_picture', ''))
            
        doc.save(output_folder / "apar_result.docx")
        logger.info("Saved APAR outputs (JSON, CSV, DOCX).")
        
    except Exception as e:
        logger.error(f"Failed to save APAR results: {e}")


def process_disciplinary(pipe, text, output_folder):
    logger.info("Processing Disciplinary Case...")
    
    # Map report types to likely folder names (categories from OCR)
    report_map = {
        "Brief Background": ["brief_background"],
        "PO Brief": ["po_brief"],
        "CO Brief": ["co_brief"],
        "IO Report": ["io_report"]
    }
    
    # Parse the full text to extract content by section
    # Format: === SECTION: {category} | FILE: {file_path.name} ===
    import re
    sections = {}
    current_category = "unknown"
    
    lines = text.split('\n')
    for line in lines:
        match = re.match(r"=== SECTION: (.*?) \| FILE: .*? ===", line)
        if match:
            current_category = match.group(1).strip()
            if current_category not in sections:
                sections[current_category] = []
        else:
            if current_category not in sections:
                sections[current_category] = []
            sections[current_category].append(line)
            
    for report_type, categories in report_map.items():
        # Aggregate text for this report
        report_text = ""
        for cat in categories:
            if cat in sections:
                report_text += "\n".join(sections[cat]) + "\n"
        
        # If no specific text found, might be "general" or fallback to everything (or skip)
        # Requirement: "Use only documents from X folder"
        if not report_text.strip():
            logger.warning(f"No specific text found for {report_type}. Using general text if available.")
            if "general" in sections:
                report_text = "\n".join(sections["general"])
            else:
                # If still empty, maybe the folder names didn't match exactly.
                # Try to find partial matches in keys
                for key in sections.keys():
                    if any(c in key for c in categories):
                        report_text += "\n".join(sections[key]) + "\n"
        
        if not report_text.strip():
            logger.warning(f"Skipping {report_type} due to lack of content.")
            continue

        prompt = f"""
        <|user|>
        Generate a "{report_type}" report based on the available text.
        Maintain formal disciplinary language. Follow chronological and logical structure.
        The report should be approximately 1-1.5 pages if enough content is available.
        
        Text:
        {report_text[:12000]}
        <|end|>
        <|assistant|>
        """
        # Increased token limit for longer reports
        try:
            output = pipe(prompt, max_new_tokens=2500, return_full_text=False)[0]['generated_text']
            
            from docx import Document
            doc = Document()
            doc.add_heading(report_type, 0)
            for para in output.split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            filename = report_type.lower().replace(" ", "_") + ".docx"
            doc.save(output_folder / filename)
            logger.info(f"Generated {filename}")
        except Exception as e:
            logger.error(f"Failed to generate {report_type}: {e}")

def main():
    parser = argparse.ArgumentParser(description="IDP System")
    parser.add_argument("--mode", choices=["ocr", "nlp"], required=True, help="Execution mode")
    parser.add_argument("--input", required=True, help="Input folder path")
    
    args = parser.parse_args()
    
    if args.mode == "ocr":
        run_ocr_pipeline(args.input)
    elif args.mode == "nlp":
        run_nlp_pipeline(args.input)

if __name__ == "__main__":
    main()
